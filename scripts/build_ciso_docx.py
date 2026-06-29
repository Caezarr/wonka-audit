from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "SECURITY-CISO.md"
OUT = ROOT / "output" / "docs" / "wonka-ai-usage-audit-ciso-security-brief.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUT)
    print(OUT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "Wonka AI Usage Audit - CISO Security Brief"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = "Local-first AI usage audit | wonka-ai.com | wonka-ai.com/cgv"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    pending_table: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                flush_table(doc, pending_table)
                pending_table.clear()
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if is_table_line(line):
            pending_table.append(line)
            i += 1
            continue
        flush_table(doc, pending_table)
        pending_table.clear()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
        elif re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^(\d+)\.\s+(.*)", stripped)
            add_numbered_item(doc, match.group(1), match.group(2))
        elif stripped.startswith("- "):
            add_inline_runs(doc.add_paragraph(style="List Bullet"), stripped[2:].strip())
        elif stripped.startswith("|"):
            pending_table.append(line)
        else:
            para = doc.add_paragraph()
            add_inline_runs(para, stripped)
        i += 1

    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))
    flush_table(doc, pending_table)


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run("Security, privacy and ISO 27001-friendly operating model")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Summary",
        "The tool runs locally, creates a local PDF and aggregate JSON export, and does not upload prompts, source code, secrets or raw conversations by default.",
    )


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f"{label}: ")
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    add_inline_runs(para, body)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=100, bottom=100, start=160, end=160)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)


def add_inline_runs(para, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(40, 40, 40)
        elif part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        else:
            para.add_run(part)


def add_numbered_item(doc: Document, number: str, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.28)
    para.paragraph_format.first_line_indent = Inches(-0.28)
    para.paragraph_format.space_after = Pt(4)
    marker = para.add_run(f"{number}. ")
    marker.font.bold = True
    add_inline_runs(para, text)


def is_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def flush_table(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, 9360)

    widths = choose_widths(cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            text = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            add_inline_runs(para, text)
            for run in para.runs:
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.bold = True
        set_row_keep_together(table.rows[r_idx])

    set_column_widths(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def choose_widths(cols: int) -> list[int]:
    if cols == 2:
        return [3000, 6360]
    if cols == 3:
        return [2600, 3380, 3380]
    return [int(9360 / cols)] * cols


def set_table_width(table, width_dxa: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_column_widths(table, widths: list[int]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                tc = row.cells[idx]._tc
                tc_pr = tc.get_or_add_tcPr()
                tc_w = tc_pr.tcW
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_row_keep_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise
